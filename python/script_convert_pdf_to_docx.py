import os


def convert_pdf_to_docx_01_pdf2docx(pdf_path, docx_path):
    from pdf2docx import Converter
    cv = Converter(pdf_path)
    cv.convert(docx_path, start=0, end=None)  # Convert all pages
    cv.close()

def convert_pdf_to_docx_02_aspose(pdf_path, docx_path):
    import aspose.pdf
    pdf_doc = aspose.pdf.Document(pdf_path)
    pdf_doc.save(docx_path, aspose.pdf.SaveFormat.DOC_X)

def convert_pdf_to_docx_03_fitz(pdf_path, docx_path):
    import fitz  # PyMuPDF
    from docx import Document
    doc = Document()
    pdf = fitz.open(pdf_path)
    for page in pdf:
        text = page.get_text("text")
        doc.add_paragraph(text)
    doc.save(docx_path)

def convert_pdf_to_docx_04_pdfminer(pdf_path, docx_path):
    from pdfminer.high_level import extract_text
    from docx import Document
    text = extract_text(pdf_path)
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(docx_path)

def convert_pdf_to_docx(pdf_path, docx_path):
    option = 1
    if option == 1:
        convert_pdf_to_docx_01_pdf2docx(pdf_path, docx_path)
    elif option == 2:
        convert_pdf_to_docx_02_aspose(pdf_path, docx_path)
    elif option == 3:
        convert_pdf_to_docx_03_fitz(pdf_path, docx_path)
    elif option == 4:
        convert_pdf_to_docx_04_pdfminer(pdf_path, docx_path)

def main():
    current_folder_full_path = os.getcwd()
    parent_folder_full_path = os.path.dirname(current_folder_full_path)
    subfolder_name = 'books'
    subfolder_full_path = os.path.join(parent_folder_full_path, subfolder_name)
    file_name = 'paramedics_full_course'
    pdf_file_name = file_name + '.pdf'
    docx_file_name = file_name + '.docx'

    pdf_file_full_path = os.path.join(subfolder_full_path, pdf_file_name)
    docx_file_full_path = os.path.join(subfolder_full_path, docx_file_name)

    convert_pdf_to_docx(pdf_file_full_path, docx_file_full_path)


main()
